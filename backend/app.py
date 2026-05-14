from fastapi import FastAPI, HTTPException, File, UploadFile, Form
import PyPDF2
import docx
from io import BytesIO
import csv
from pydantic import BaseModel
from typing import Optional, List
import uuid
from datetime import datetime

from llm_client import LLMClient
from prompts import (
    GENERATE_QUIZ_PROMPT,
    GENERATE_QUIZ_FROM_FILE_PROMPT,
    REGENERATE_QUIZ_PROMPT,
    SUMMARIZE_FOR_QUIZ_PROMPT
)
from db_client import save_generation_log

from logger_config import setup_logger

logger = setup_logger(__name__)

app = FastAPI(title="Test Generation Service")

llm_client = LLMClient()


class GenerateRequest(BaseModel):
    topic: str
    num_questions: int = 5
    model: str = "local"  # Добавляем выбор модели


class GenerateResponse(BaseModel):
    request_id: str
    topic: str
    num_questions: int
    quiz_text: str
    timestamp: str
    used_model: str


class ModelInfo(BaseModel):
    id: str
    name: str
    available: bool


class RegenerateRequest(BaseModel):
    topic: str
    num_questions: int
    old_quiz_text: str
    model: str = "local"


@app.post("/regenerate")
async def regenerate_quiz(request: RegenerateRequest):
    """Перегенерирует тест, исключая вопросы из предыдущего теста"""
    request_id = str(uuid.uuid4())
    timestamp = datetime.now().isoformat()

    logger.info(f"Перегенерация: model={request.model}, topic={request.topic}, num_questions={request.num_questions}")

    # Промпт с исключением старых вопросов
    prompt = REGENERATE_QUIZ_PROMPT.format(
        num_questions=request.num_questions,
        topic=request.topic,
        old_quiz_text=request.old_quiz_text
    )
    logger.debug(f"Промпт для перегенерации (первые 200 символов): {prompt[:200]}...")

    # Переключаем модель, если нужно
    if request.model != llm_client.current_model:
        logger.debug(f"Переключение модели с {llm_client.current_model} на {request.model}")
        success = llm_client.switch_model(request.model)
        if not success:
            logger.error(f"Не удалось переключить модель на {request.model}")
            raise HTTPException(status_code=400, detail=f"Модель {request.model} недоступна")

    try:
        quiz_text = llm_client.generate(prompt)
        logger.info(f"Перегенерация успешно завершена, длина текста: {len(quiz_text)} символов")
    except Exception as e:
        logger.error(f"Ошибка при перегенерации: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

    # Сохраняем лог
    save_generation_log(request_id, request.topic, request.num_questions, quiz_text)
    logger.debug(f"Лог сохранён: request_id={request_id}")

    return {
        "request_id": request_id,
        "topic": request.topic,
        "num_questions": request.num_questions,
        "quiz_text": quiz_text,
        "timestamp": timestamp,
        "used_model": llm_client.current_model
    }


@app.post("/generate", response_model=GenerateResponse)
async def generate_quiz(request: GenerateRequest):
    request_id = str(uuid.uuid4())
    timestamp = datetime.now().isoformat()

    logger.info(f"Генерация теста: model={request.model}, topic={request.topic}, num_questions={request.num_questions}")

    # Переключаем модель, если нужно
    if request.model != llm_client.current_model:
        logger.debug(f"Переключение модели с {llm_client.current_model} на {request.model}")
        llm_client.switch_model(request.model)

    prompt = GENERATE_QUIZ_PROMPT.format(
        topic=request.topic,
        num_questions=request.num_questions
    )
    logger.debug(f"Промпт для генерации (первые 200 символов): {prompt[:200]}...")

    try:
        quiz_text = llm_client.generate(prompt)
        logger.info(f"Генерация успешно завершена, длина текста: {len(quiz_text)} символов")
    except Exception as e:
        logger.error(f"Ошибка при генерации: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

    # Сохраняем лог в БД
    save_generation_log(request_id, request.topic, request.num_questions, quiz_text)
    logger.debug(f"Лог сохранён: request_id={request_id}")

    return GenerateResponse(
        request_id=request_id,
        topic=request.topic,
        num_questions=request.num_questions,
        quiz_text=quiz_text,
        timestamp=timestamp,
        used_model=llm_client.current_model
    )


@app.post("/generate-from-file")
async def generate_quiz_from_file(
        file: UploadFile = File(...),
        num_questions: int = Form(5),
        model: str = Form("local")
):
    """Генерирует тест из содержимого загруженного файла"""
    request_id = str(uuid.uuid4())
    timestamp = datetime.now().isoformat()

    logger.info(f"Генерация из файла: model={model}, num_questions={num_questions}, filename={file.filename}")

    # Извлекаем текст из файла
    try:
        file_content = await extract_text_from_file(file)
        logger.info(f"Текст из файла извлечён, длина: {len(file_content)} символов")
    except Exception as e:
        logger.error(f"Ошибка при извлечении текста из файла {file.filename}: {e}", exc_info=True)
        raise HTTPException(status_code=400, detail=f"Ошибка при обработке файла: {str(e)}")

    if not file_content or len(file_content.strip()) < 50:
        logger.warning(
            f"Файл {file.filename} слишком маленький или пустой: {len(file_content) if file_content else 0} символов")
        raise HTTPException(status_code=400, detail="Файл слишком маленький или пустой")

    prompt = GENERATE_QUIZ_FROM_FILE_PROMPT.format(
        num_questions=num_questions,
        file_content=file_content[:5000]  # Логируем только начало
    )
    logger.debug(f"Промпт для генерации из файла (первые 500 символов): {prompt[:500]}...")

    # Переключаем модель, если нужно
    if model != llm_client.current_model:
        logger.debug(f"Переключение модели с {llm_client.current_model} на {model}")
        success = llm_client.switch_model(model)
        if not success:
            logger.error(f"Не удалось переключить модель на {model}")
            raise HTTPException(status_code=400, detail=f"Модель {model} недоступна")

    try:
        quiz_text = llm_client.generate(prompt)
        logger.info(f"Генерация из файла успешно завершена, длина текста: {len(quiz_text)} символов")
    except Exception as e:
        logger.error(f"Ошибка при генерации из файла: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

    # Сохраняем лог
    save_generation_log(request_id, f"Из файла: {file.filename}", num_questions, quiz_text)
    logger.debug(f"Лог сохранён: request_id={request_id}, topic=Из файла: {file.filename}")

    return {
        "request_id": request_id,
        "topic": f"Из файла: {file.filename}",
        "num_questions": num_questions,
        "quiz_text": quiz_text,
        "timestamp": timestamp,
        "used_model": llm_client.current_model
    }


async def extract_text_from_file(file: UploadFile) -> str:
    """Извлекает текст из разных типов файлов"""
    content = await file.read()
    logger.debug(f"Извлечение текста из файла: {file.filename}, тип: {file.content_type}")

    if file.filename.endswith('.txt'):
        text = content.decode('utf-8')
        logger.debug(f"TXT файл прочитан, длина: {len(text)} символов")
        return text

    elif file.filename.endswith('.pdf'):
        pdf_file = BytesIO(content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for i, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            text += page_text
            logger.debug(f"PDF страница {i + 1}/{len(pdf_reader.pages)} извлечена, длина: {len(page_text)} символов")
        logger.info(f"PDF файл обработан, всего страниц: {len(pdf_reader.pages)}, итоговая длина: {len(text)} символов")
        return text

    elif file.filename.endswith('.docx'):
        doc = docx.Document(BytesIO(content))
        text = ""
        for i, paragraph in enumerate(doc.paragraphs):
            text += paragraph.text + "\n"
        logger.debug(
            f"DOCX файл обработан, всего параграфов: {len(doc.paragraphs)}, длина текста: {len(text)} символов")
        return text

    elif file.filename.endswith('.csv'):
        csv_text = content.decode('utf-8')
        logger.debug(f"CSV файл прочитан, длина: {len(csv_text)} символов")
        return csv_text

    else:
        logger.warning(f"Неподдерживаемый тип файла: {file.filename}")
        raise HTTPException(status_code=400, detail=f"Неподдерживаемый тип файла: {file.filename}")


@app.get("/models")
async def get_available_models():
    """Эндпоинт для получения списка доступных моделей"""
    models = llm_client.get_available_models()
    logger.debug(f"Запрос списка моделей, возвращено {len(models)} моделей")
    return {"models": models}


@app.get("/health")
async def health():
    """Эндпоинт для проверки здоровья сервиса"""
    status = {
        "status": "ok",
        "gemini_available": llm_client.gemini_available,
        "ollama_available": llm_client.ollama_available,
        "current_model": llm_client.current_model
    }
    logger.debug(f"Health check: {status}")
    return status


if __name__ == "__main__":
    import uvicorn

    logger.info("🚀 Запуск FastAPI сервера на порту 8000")
    uvicorn.run(app, host="0.0.0.0", port=8000)